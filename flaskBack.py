from flask import Flask, request, jsonify, send_file, render_template
import pdfplumber
from groq import Groq
from docx import Document
from fpdf import FPDF
from io import BytesIO

app = Flask(__name__)

# Initialize Groq client with your API key
API_KEY = "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"
client = Groq(api_key=API_KEY)

# Function to split text into chunks of a maximum token length
def chunk_text(text, max_tokens=1024):
    words = text.split()
    chunks = []
    chunk = []

    for word in words:
        chunk.append(word)
        if len(" ".join(chunk)) > max_tokens:
            chunks.append(" ".join(chunk[:-1]))
            chunk = [chunk[-1]]

    if chunk:
        chunks.append(" ".join(chunk))

    return chunks

@app.route("/")
def home():
    return render_template("index.html")  # Serve the HTML UI

@app.route("/upload", methods=["POST"])
def upload_pdf():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    uploaded_file = request.files["file"]

    try:
        # Extract text from the uploaded PDF
        with pdfplumber.open(uploaded_file) as pdf:
            full_text = ""
            for page in pdf.pages:
                full_text += page.extract_text()

        return jsonify({"text": full_text}), 200
    except Exception as e:
        return jsonify({"error": f"Failed to process PDF: {str(e)}"}), 500

@app.route("/generate_quiz", methods=["POST"])
def generate_quiz():
    data = request.json
    full_text = data.get("text", "")

    if not full_text:
        return jsonify({"error": "No text provided"}), 400

    try:
        text_chunks = chunk_text(full_text, max_tokens=1024)
        final_quiz = ""

        for chunk in text_chunks:
            completion = client.chat.completions.create(
                model="llama3-8b-8192",
                messages=[{
                    "role": "user",
                    "content": f"Generate subjective questions that assess thorough understanding of the following text:\n{chunk}"
                }],
                temperature=0.7,
                max_tokens=1024,
                top_p=1,
                stream=True,
                stop=None,
            )

            quiz_text = ""
            for response_chunk in completion:
                quiz_text += response_chunk.choices[0].delta.content or ""

            final_quiz += quiz_text.strip() + "\n\n"

        return jsonify({"quiz": final_quiz.strip()}), 200
    except Exception as e:
        return jsonify({"error": f"Failed to generate quiz: {str(e)}"}), 500

@app.route("/download", methods=["POST"])
def download_question_paper():
    data = request.json
    questions = data.get("questions", "").split("\n")

    if not questions:
        return jsonify({"error": "No questions provided"}), 400

    try:
        # Generate DOCX file
        doc = Document()
        doc.add_heading("Generated Question Paper", 0)
        doc.add_heading("Instructions:", level=1)
        doc.add_paragraph("Answer all the questions based on the content of the PDF.")
        doc.add_heading("Questions:", level=1)
        for question in questions:
            doc.add_paragraph(question.strip())

        doc_buffer = BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)

        return send_file(
            doc_buffer,
            as_attachment=True,
            download_name="question_paper.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except Exception as e:
        return jsonify({"error": f"Failed to generate document: {str(e)}"}), 500

if __name__ == "__main__":
    app.run()
